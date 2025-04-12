import os
from typing import Optional, Dict
from pydantic import BaseModel, Field
import enum
from docx import Document

from langchain.tools import BaseTool
from langchain.callbacks.manager import (
    CallbackManagerForToolRun,
    AsyncCallbackManagerForToolRun,
)


def fill_docx_template(template_path, data, output_path):
    """
    Fills a Word document (.docx) template with data and saves the filled document to a new file.

    Args:
        template_path (str): The path to the Word document template file.
        data (dict): A dictionary where keys are the placeholders in the template
                     (e.g., "{{name}}", "{{age}}") and values are the data to replace them with.
        output_path (str): The path to save the filled Word document.
    """

    try:
        document = Document(template_path)

        # Iterate through paragraphs and replace placeholders
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                for placeholder, value in data.items():
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))  # Convert value to string

        # Iterate through tables and replace placeholders
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in data.items():
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))  # Convert value to string

        document.save(output_path)
        print(f"Document saved to: {output_path}")

    except FileNotFoundError:
        print(f"Error: Template file not found at {template_path}")
    except Exception as e:
        print(f"An error occurred: {e}")


def format_tien(so_tien: int) -> str:
    """
    Định dạng số tiền từ kiểu int sang string, thêm dấu chấm mỗi 3 chữ số.

    Ví dụ:
    format_tien(1000) == "1.000"
    format_tien(1000000) == "1.000.000"
    format_tien(100) == "100"
    """

    so_tien_str = str(so_tien)
    ket_qua = ""
    count = 0

    for i in range(len(so_tien_str) - 1, -1, -1):
        ket_qua = so_tien_str[i] + ket_qua
        count += 1
        if count == 3 and i != 0:
            ket_qua = "." + ket_qua
            count = 0
    return ket_qua


class KinhPhiType(str, enum.Enum):
    """Loại kinh phí"""

    MUA_SAM = "xin mua sắm thiết bị"
    OTHER = "loại khác"


class ToTrinhKinhPhiDetails(BaseModel):
    """Input schema cho tờ trình xin kinh phí"""

    tieu_de: Optional[str] = Field(..., description="Tiêu đề của tờ trình. Có thể tự tạo dựa trên loại kinh phí.")
    loai_kinh_phi: KinhPhiType = Field(
        ..., description="Loại kinh phí: xin mua sắm thiết bị hoặc loại khác. Bắt buộc phải do người dùng cung cấp."
    )
    ngay_lap: Optional[str] = Field(
        ..., description="Ngày lập tờ trình (ISO 8601). Có thể tự động tạo là ngày hôm nay."
    )
    nguoi_de_xuat: str = Field(..., description="Người đề xuất (Họ và tên).")
    don_vi: str = Field(..., description="Đơn vị công tác.")
    noi_dung: str = Field(
        ...,
        description="Nội dung chi tiết của đề xuất. Người dùng bắt buộc cần cung cấp lý do đề xuất. Sau đó phần nội dung có thể tự được sinh ra với độ dài khoảng 5 câu.",
    )
    tong_kinh_phi: float = Field(..., description="Tổng kinh phí đề xuất (số). Bắt buộc phải do người dùng cung cấp.")
    don_vi_tien_te: Optional[str] = Field(default="VND", description="Đơn vị tiền tệ (mặc định là VND).")
    de_xuat_khac: Optional[str] = Field(description="Các đề xuất khác (nếu có).", default=None)


class TaoToTrinhKinhPhiTool(BaseTool):
    name: str = "tao_to_trinh_kinh_phi"
    description: str = "Tạo một tờ trình xin kinh phí với các chi tiết được cung cấp."
    args_schema: type[BaseModel] = ToTrinhKinhPhiDetails
    return_direct: bool = False

    def _run(
        self,
        tieu_de: Optional[str],
        loai_kinh_phi: KinhPhiType,
        ngay_lap: Optional[str],
        nguoi_de_xuat: str,
        don_vi: str,
        noi_dung: Optional[str],
        tong_kinh_phi: float,
        don_vi_tien_te: Optional[str] = "VND",
        de_xuat_khac: Optional[str] = None,
        run_manager: Optional[CallbackManagerForToolRun] = None,
    ) -> Dict:
        """Tạo tờ trình xin kinh phí. Hãy gọi công cụ này khi có thể, nếu thiếu trường nào công cụ sẽ báo lại.

        Args:
            tieu_de (Optional[str], optional): Tiêu đề của tờ trình. Có thể tự tạo dựa trên loại kinh phí.
            loai_kinh_phi (KinhPhiType): Loại kinh phí: xin mua sắm thiết bị hoặc loại khác. Bắt buộc phải do người dùng cung cấp.
            ngay_lap (Optional[str], optional): Ngày lập tờ trình (ISO 8601). Có thể tự động tạo là ngày hôm nay.
            nguoi_de_xuat (str): Người đề xuất (Họ và tên).
            don_vi (str): Đơn vị công tác.
            noi_dung (str): Nội dung chi tiết của đề xuất. Người dùng bắt buộc cần cung cấp lý do đề xuất. Sau đó phần nội dung có thể tự được sinh ra với độ dài khoảng 5 câu.
            tong_kinh_phi (float): Tổng kinh phí đề xuất. Bắt buộc phải do người dùng cung cấp.
            don_vi_tien_te (Optional[str], optional): Đơn vị tiền tệ. Mặc định là "VND".
            de_xuat_khac (Optional[str], optional): Các đề xuất khác. Mặc định là None.
            run_manager (Optional[CallbackManagerForToolRun], optional): Quản lý callback. Mặc định là None.

        Returns:
            Dict: Kết quả tạo tờ trình.
                - Nếu thành công: {"status": "success", "message": "Tờ trình đã được tạo thành công.", "data": { <thông tin tờ trình> } }
                - Nếu lỗi: {"status": "error", "message": ...}
        """
        try:
            tong_kinh_phi = format_tien(int(tong_kinh_phi))

            to_trinh_data = {
                "tieu_de": tieu_de,
                "loai_kinh_phi": loai_kinh_phi.value,
                "ngay_lap": ngay_lap,
                "nguoi_de_xuat": nguoi_de_xuat,
                "don_vi": don_vi,
                "noi_dung": noi_dung,
                "tong_kinh_phi": tong_kinh_phi,
                "don_vi_tien_te": don_vi_tien_te,
                "de_xuat_khac": de_xuat_khac,
            }

            if loai_kinh_phi == KinhPhiType.MUA_SAM:
                template_path = "./assets/template_to_trinh_xin_mua_sam_thiet_bi.docx"
                os.makedirs("output", exist_ok=True)
                fill_docx_template(
                    template_path,
                    {
                        "{{TIEU_DE}}": tieu_de.upper(),
                        "{{loai_kinh_phi}}": loai_kinh_phi,
                        "{{ngay}}": ngay_lap.split("-")[2],
                        "{{thang}}": ngay_lap.split("-")[1],
                        "{{nam}}": ngay_lap.split("-")[0],
                        "{{nguoi_de_xuat}}": nguoi_de_xuat,
                        "{{NGUOI_DE_XUAT}}": nguoi_de_xuat.upper(),
                        "{{don_vi}}": don_vi,
                        "{{noi_dung}}": noi_dung,
                        "{{tong_kinh_phi}}": tong_kinh_phi,
                        "{{don_vi_tien_te}}": don_vi_tien_te,
                        "{{de_xuat_khac}}": "" if de_xuat_khac is None else de_xuat_khac,
                    },
                    f"./output/{run_manager.run_id}.docx",
                )
                to_trinh_data["docx"] = f"./output/{run_manager.run_id}.docx"
            elif loai_kinh_phi == KinhPhiType.OTHER:
                raise Exception("Loại tờ trình này chưa được hỗ trợ.")
            else:
                raise Exception("Loại tờ trình này chưa được hỗ trợ.")

            return {
                "status": "success",
                "message": "Tờ trình đã được tạo thành công.",
                "data": to_trinh_data,
            }

        except Exception as e:
            return {"status": "error", "message": str(e)}

    async def _arun(
        self,
        tieu_de: Optional[str],
        loai_kinh_phi: KinhPhiType,
        ngay_lap: Optional[str],
        nguoi_de_xuat: str,
        don_vi: str,
        noi_dung: Optional[str],
        tong_kinh_phi: float,
        don_vi_tien_te: Optional[str] = "VND",
        de_xuat_khac: Optional[str] = None,
        run_manager: Optional[AsyncCallbackManagerForToolRun] = None,
    ) -> Dict:
        """Asynchronous execution for creating the to trinh kinh phi."""
        return self._run(
            tieu_de,
            loai_kinh_phi,
            ngay_lap,
            nguoi_de_xuat,
            don_vi,
            noi_dung,
            tong_kinh_phi,
            don_vi_tien_te,
            de_xuat_khac,
            run_manager=run_manager.get_sync(),
        )
